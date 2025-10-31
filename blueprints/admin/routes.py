from flask import Blueprint, render_template, redirect, jsonify , url_for, flash, request, current_app, send_file
from flask_login import login_required, current_user
from models.user import User
from models.booking import Booking
from models.room import Room
from models.application_rooms import ApplicationRoom
from models.payment import Payment
from models.service_request import ServiceRequest
from models.complain import Complain
from models.complain_image import ComplainImage
from models.announcement import Announcement 
import os, time, io 
from models.room_image import RoomImage
from models.notification import Notification
from extensions import db
from datetime import datetime
from werkzeug.utils import secure_filename
from sqlalchemy import or_, func, extract
import unicodedata, calendar
import pandas as pd
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from math import ceil

admin_bp = Blueprint("admin", __name__, template_folder="../../templates/admin")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in current_app.config["ALLOWED_EXTENSIONS"]

# Middleware: chỉ admin mới vào được
@admin_bp.before_request
def restrict_to_admin():
    if not current_user.is_authenticated or current_user.role != "admin":
        flash("Bạn không có quyền truy cập trang quản trị.", "danger")
        return redirect(url_for("auth.login"))
#-------------------------------------------------------
#Dashboard
@admin_bp.route("/dashboard")
@login_required
def dashboard():
    # Lấy tháng, năm từ query params (mặc định là tháng/năm hiện tại)
    selected_month = request.args.get('month', type=int)
    selected_year = request.args.get('year', type=int)
    
    # Nếu không có giá trị, dùng tháng/năm hiện tại
    if not selected_month:
        selected_month = datetime.now().month
    if not selected_year:
        selected_year = datetime.now().year

    # Tổng sinh viên, phòng, hóa đơn chưa thanh toán
    total_students = User.query.filter_by(role="student").count()
    total_rooms = Room.query.count()
    unpaid_bills = Payment.query.filter(Payment.status != "success").count()

    # Thống kê doanh thu theo loại dịch vụ trong tháng đã chọn
    revenue_by_service = (
        db.session.query(Payment.service_type, func.sum(Payment.amount))
        .filter(
            Payment.month_paid == selected_month,
            Payment.year_paid == selected_year,
            Payment.status == "success"
        )
        .group_by(Payment.service_type)
        .all()
    )
    service_name_map = {
    "deposit": "Tiền phòng",
    "utilities": "Điện & Nước",
    "trash": "Thu gom rác",
    "maintenance": "Bảo trì - Sửa chữa"
    }
    service_labels = [service_name_map.get(r[0], "Khác") for r in revenue_by_service]
    service_values = [float(r[1]) if r[1] else 0 for r in revenue_by_service]

    # Tổng tiền hóa đơn tháng đã chọn
    total_revenue = sum(service_values)

    # Số sinh viên đăng ký mới theo tuần trong tháng đã chọn
    start_of_month = datetime(selected_year, selected_month, 1)
    last_day = calendar.monthrange(selected_year, selected_month)[1]
    end_of_month = datetime(selected_year, selected_month, last_day, 23, 59, 59)
    
    # Lấy tất cả booking trong tháng
    bookings_this_month = Booking.query.filter(
        Booking.created_at >= start_of_month,
        Booking.created_at <= end_of_month
    ).order_by(Booking.created_at).all()
    
    # Chia theo tuần 
    week_counts = {}
    for booking in bookings_this_month:
        # Tính tuần thứ mấy trong tháng (bắt đầu từ ngày 1)
        day_of_month = booking.created_at.day
        week_number = ((day_of_month - 1) // 7) + 1
        week_label = f"Tuần {week_number}"
        week_counts[week_label] = week_counts.get(week_label, 0) + 1
    
    # Đảm bảo có đủ 4-5 tuần
    max_weeks = (last_day // 7) + (1 if last_day % 7 else 0)
    reg_labels = [f"Tuần {i+1}" for i in range(max_weeks)]
    reg_values = [week_counts.get(label, 0) for label in reg_labels]

    # Phòng trống vs đã có người
    occupied = Room.query.filter(Room.available < Room.capacity).count()
    empty = Room.query.filter(Room.available == Room.capacity).count()

    # Danh sách tháng và năm
    months = list(range(1, 13))
    years = list(range(2020, datetime.now().year + 2))

    return render_template(
        "admin/admin_dashboard.html",
        total_students=total_students,
        total_rooms=total_rooms,
        unpaid_bills=unpaid_bills,
        total_revenue=total_revenue,
        service_labels=service_labels,
        service_values=service_values,
        reg_labels=reg_labels,
        reg_values=reg_values,
        occupied=occupied,
        empty=empty,
        selected_month=selected_month,
        selected_year=selected_year,
        months=months,
        years=years
    )

# Xuất file excel thống kê
@admin_bp.route("/dashboard/export_excel")
@login_required
def export_dashboard_excel():
    total_students = User.query.filter_by(role="student").count()
    total_rooms = Room.query.count()
    total_bookings = Booking.query.count()
    total_payments = Payment.query.count()
    unpaid_payments = Payment.query.filter_by(status="pending").count()

    df_summary = pd.DataFrame([
        {"Hạng mục": "Tổng sinh viên", "Số lượng": total_students},
        {"Hạng mục": "Tổng phòng", "Số lượng": total_rooms},
        {"Hạng mục": "Đang thuê phòng", "Số lượng": total_bookings},
        {"Hạng mục": "Tổng hóa đơn", "Số lượng": total_payments},
        {"Hạng mục": "Chưa thanh toán", "Số lượng": unpaid_payments},
    ])

    # Chi tiết hóa đơn
    payments = Payment.query.join(User, Payment.user_id == User.id).join(Room, Payment.room_id == Room.id).all()
    payment_data = []
    for p in payments:
        payment_data.append({
            "MSSV": p.user.student_id if hasattr(p.user, "student_id") else "",
            "Họ tên": p.user.fullname if hasattr(p.user, "fullname") else "",
            "Phòng": f"{p.room.room_number} ({p.room.block})" if hasattr(p.room, "room_number") else "",
            "Loại dịch vụ": {
                "rent": "Tiền phòng",
                "deposit": "Tiền đặt cọc",
                "utilities": "Điện & Nước",
                "maintenance": "Bảo trì",
                "trash": "Thu gom rác"
            }.get(p.service_type, p.service_type),
            "Tháng": p.month_paid,
            "Năm": p.year_paid,
            "Số tiền (VNĐ)": float(p.amount or 0),
            "Trạng thái": "Đã thanh toán" if p.status == "success" else "Chưa thanh toán",
            "Phương thức": p.payment_method or "",
            "Ngày tạo": p.created_at.strftime("%d/%m/%Y") if p.created_at else "",
            "Cập nhật cuối": p.updated_at.strftime("%d/%m/%Y") if p.updated_at else "",
        })
    df_payments = pd.DataFrame(payment_data)

    #Danh sách sinh viên
    students = User.query.filter_by(role="student").all()
    student_data = [{
        "MSSV": s.student_id,
        "Họ tên": s.fullname,
        "Lớp": s.class_id,
        "Email": s.email,
        "SĐT": s.phone_number,
        "Ngày tạo": s.created_at.strftime("%d/%m/%Y") if s.created_at else ""
    } for s in students]
    df_students = pd.DataFrame(student_data)

    #Ghi ra file Excel
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_summary.to_excel(writer, index=False, sheet_name="Tổng quan")
        df_students.to_excel(writer, index=False, sheet_name="Sinh viên")
        df_payments.to_excel(writer, index=False, sheet_name="Hóa đơn chi tiết")

        workbook = writer.book
        worksheet = writer.sheets["Tổng quan"]
        header_format = workbook.add_format({"bold": True, "bg_color": "#CCE5FF", "border": 1})
        for col_num, value in enumerate(df_summary.columns.values):
            worksheet.write(0, col_num, value, header_format)

    output.seek(0)
    filename = f"Thong_ke_KTX_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.xlsx"

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

#-------------------------------------------------------
# Quản lý phòng
@admin_bp.route("/rooms")
@login_required
def manage_rooms():
    address = request.args.get("address")
    block = request.args.get("block")
    room_number = request.args.get("room_number")

    query = Room.query

    if address:
        query = query.filter_by(address=address)
    if block:
        query = query.filter_by(block=block)
    if room_number:
        query = query.filter_by(room_number=room_number)

    rooms = query.all()

    return render_template("admin/manage_rooms.html", rooms=rooms)

# Chi tiết phòng
@admin_bp.route("/rooms/<int:room_id>")
@login_required
def room_detail(room_id):
    room = Room.query.get_or_404(room_id)

    bookings = (
        Booking.query
        .filter_by(room_id=room.id)
        .join(User, User.id == Booking.user_id)
        .add_entity(User)
        .all()
    )

    active_bookings = []
    for booking, user in bookings:
        booking.user = user
        if booking.status == "active":
            active_bookings.append(booking)

    return render_template("admin/room_detail.html", room=room, bookings=active_bookings)


# Sửa phòng
@admin_bp.route("/rooms/edit", methods=["GET", "POST"])
@admin_bp.route("/rooms/edit/<int:room_id>", methods=["GET", "POST"])
@login_required
def edit_room(room_id=None):
    room = Room.query.get(room_id) if room_id else None

    if request.method == "POST":
        block        = (request.form.get("block") or "").strip()
        room_number  = (request.form.get("room_number") or "").strip()
        address      = (request.form.get("address") or "").strip()
        capacity     = int(request.form.get("capacity") or 0)

        if room:
            room.block = block
            room.room_number = room_number
            room.address = address
            room.capacity = capacity
            flash("Cập nhật phòng thành công!", "success")
        else:  
            room = Room(
                block=block,
                room_number=room_number,
                address=address,
                capacity=capacity,
                current_occupancy=0,
            )
            db.session.add(room)
            db.session.flush()

            flash("Thêm phòng thành công!", "success")

        files = request.files.getlist("images")
        if files:
            base_folder = os.path.join(current_app.config["UPLOAD_FOLDER_ROOMS"], address, block)
            os.makedirs(base_folder, exist_ok=True)

            for f in files:
                if not (f and f.filename and allowed_file(f.filename)):
                    continue
                safe = secure_filename(f.filename)
                filename = f"{int(time.time()*1000)}_{safe}"
                save_path = os.path.join(base_folder, filename)
                f.save(save_path)

                rel_url = f"/static/img/room/{address}/{block}/{filename}"
                db.session.add(RoomImage(room_id=room.id, image_url=rel_url))

        db.session.commit()
        return redirect(url_for("admin.room_detail", room_id=room.id))

    return render_template("admin/edit_room.html", room=room)

# Xoá sinh viên khỏi phòng (kick)
@admin_bp.route("/rooms/<int:room_id>/kick/<int:booking_id>", methods=["POST"])
@login_required
def kick_student(room_id, booking_id):
    booking = Booking.query.get_or_404(booking_id)

    if booking.status != "active":
        flash("Sinh viên này không còn ở phòng.", "warning")
        return redirect(url_for("admin.room_detail", room_id=room_id))

    # Cập nhật trạng thái booking
    booking.status = "canceled"
    booking.end_date = db.func.current_date()

    # Cập nhật phòng
    room = booking.room
    if room:
        room.available = room.available + 1

    db.session.commit()

    # Tạo thông báo cho sinh viên
    notif = Notification(
        user_id=booking.user_id,
        message=f"Bạn đã bị kích khỏi phòng {room.room_number}."
    )
    db.session.add(notif)
    db.session.commit()
    

    flash("Đã kick sinh viên ra khỏi phòng!", "success")
    return redirect(url_for("admin.room_detail", room_id=room_id))
#-------------------------------------------------------
# Quản lý đơn đăng ký
@admin_bp.route("/manage_applications")
@login_required
def manage_applications():
    if current_user.role != "admin":
        flash("Bạn không có quyền truy cập!", "danger")
        return redirect(url_for("home"))

    status = request.args.get("status")
    address = request.args.get("address")
    block = request.args.get("block")
    room_number = request.args.get("room_number")
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int, default=datetime.now().year)
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    query = ApplicationRoom.query.join(Room, ApplicationRoom.room_id == Room.id)

    if status:
        query = query.filter(ApplicationRoom.status == status)
    if address:
        query = query.filter(Room.address == address)
    if block:
        query = query.filter(Room.block == block)
    if room_number:
        query = query.filter(Room.room_number == room_number)
    if month:
        query = query.filter(db.extract("month", ApplicationRoom.created_at) == month)
    if year:
        query = query.filter(db.extract("year", ApplicationRoom.created_at) == year)
    if date_from:
        try:
            start_date = datetime.strptime(date_from, "%Y-%m-%d")
            query = query.filter(ApplicationRoom.created_at >= start_date)
        except:
            pass
    if date_to:
        try:
            end_date = datetime.strptime(date_to, "%Y-%m-%d")
            query = query.filter(ApplicationRoom.created_at <= end_date)
        except:
            pass

    applications = query.order_by(ApplicationRoom.created_at.desc()).all()

    return render_template(
        "admin/manage_applications.html",
        applications=applications,
        filters={
            "status": status,
            "address": address,
            "block": block,
            "room_number": room_number,
            "month": month,
            "year": year,
            "date_from": date_from,
            "date_to": date_to
        }
    )


# API: Lấy chi tiết 1 đơn (JSON) -> dùng AJAX gọi
@admin_bp.route("/api/application/<int:app_id>")
@login_required
def get_application_detail(app_id):
    if current_user.role != "admin":
        return jsonify({"error": "Forbidden"}), 403

    app = ApplicationRoom.query.get_or_404(app_id)
    room = Room.query.get(app.room_id)

    data = {
        "id": app.id,
        "fullname": app.fullname,
        "student_id": app.student_id,
        "class_id": app.class_id,
        "citizen_id": app.citizen_id,
        "email": app.email,
        "phone_number": app.phone_number,
        "room": f"{room.block}{room.room_number} - {room.address}",
        "status": app.status,
        "created_at": app.created_at.strftime("%d/%m/%Y %H:%M"),
        "can_approve": app.status == "pending",
        "can_reject": app.status == "pending",
    }

    # Thêm chi tiết từ ApplicationRoom
    data.update({
        "address": app.address,
        "relative1_name": app.relative1_name,
        "relative1_phone": app.relative1_phone,
        "relative1_birthyear": app.relative1_birthyear,
        "relative2_name": app.relative2_name,
        "relative2_phone": app.relative2_phone,
        "relative2_birthyear": app.relative2_birthyear,
        "policy_type": app.policy_type,
        "policy_proof": app.policy_proof,
        "student_photo": app.student_photo,
        "citizen_proof": app.citizen_proof,
    })

    return jsonify(data)

# Duyệt đơn
@admin_bp.route("/application/<int:app_id>/approve")
@login_required
def approve_application(app_id):
    if current_user.role != "admin":
        flash("Bạn không có quyền truy cập!", "danger")
        return redirect(url_for("home"))

    application = ApplicationRoom.query.get_or_404(app_id)
    room = Room.query.get(application.room_id)

    # Kiểm tra phòng còn chỗ không
    if room.available <= 0:
        flash("Phòng đã hết chỗ, không thể duyệt đơn!", "danger")
        return redirect(url_for("admin.manage_applications"))

    # Kiểm tra đơn đã được xử lý chưa
    if application.status != "pending":
        flash("Đơn này đã được xử lý trước đó.", "info")
        return redirect(url_for("admin.manage_applications"))
    
    # Chỉ thay đổi trạng thái đơn, KHÔNG trừ available ở đây
    application.status = "approved_pending_deposit"

    # Tạo hóa đơn tiền cọc

    now = datetime.now()
    payment = Payment(
        application_id=application.id,
        booking_id=None,   # chưa có booking chính thức
        user_id=application.user_id,
        room_id=application.room_id,
        fullname=application.fullname,
        student_id=application.student_id,
        class_id=application.class_id,
        citizen_id=application.citizen_id,
        email=application.email,
        phone_number=application.phone_number,
        service_type="deposit",
        month_paid=datetime.now().month,
        year_paid=datetime.now().year,
        address=room.address,
        block=room.block,
        room_number=room.room_number,
        amount=room.deposit,
        payment_method="cash",
        status="pending"   
    )

    db.session.add(payment)
    db.session.commit()

    # Tạo thông báo cho sinh viên
    notif = Notification(
    user_id=application.user_id,
    message=f"Đơn đăng ký phòng {room.room_number} của bạn đã được duyệt. Vui lòng thanh toán tiền cọc để hoàn tất."
    )
    db.session.add(notif)
    db.session.commit()

    flash("Đơn đã được duyệt. Sinh viên cần thanh toán cọc trước khi được xếp phòng.", "success")
    return redirect(url_for("admin.manage_applications"))


#Từ chối đơn
@admin_bp.route("/application/<int:app_id>/reject")
@login_required
def reject_application(app_id):
    if current_user.role != "admin":
        flash("Bạn không có quyền truy cập!", "danger")
        return redirect(url_for("home"))

    application = ApplicationRoom.query.get_or_404(app_id)
    if application.status != "pending":
        flash("Đơn này đã được xử lý trước đó.", "info")
        return redirect(url_for("admin.manage_applications"))

    application.status = "rejected"
    db.session.commit()

    # Tạo thông báo cho sinh viên
    notif = Notification(
    user_id=application.user_id,
    message=f"Đơn đăng ký phòng {application.room.room_number} của bạn đã bị TỪ CHỐI."
    )
    db.session.add(notif)
    db.session.commit()
    

    flash("Đơn đã bị từ chối.", "info")
    return redirect(url_for("admin.manage_applications"))

#Xuất file docx
@admin_bp.route('/export_application/<int:application_id>')
def export_application(application_id):
    app = ApplicationRoom.query.get_or_404(application_id)
    user = app.user

    doc = Document()
    doc.add_heading("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 0)
    doc.add_paragraph("Độc lập - Tự do - Hạnh phúc", style="Intense Quote")
    doc.add_paragraph("\n\n", style=None)
    doc.add_heading("ĐƠN ĐĂNG KÝ Ở KÝ TÚC XÁ", level=0).alignment = 1

    # Thông tin cá nhân
    doc.add_heading("1. Thông tin cá nhân", level=1)
    info = [
        ("Họ và tên", user.fullname),
        ("MSSV", user.student_id),
        ("Lớp", user.class_id),
        ("CCCD", user.citizen_id),
    ]
    for label, value in info:
        doc.add_paragraph(f"{label}: {value or ''}")

    # Liên hệ
    doc.add_heading("2. Thông tin liên hệ", level=1)
    doc.add_paragraph(f"Email: {user.email or ''}")
    doc.add_paragraph(f"Số điện thoại: {user.phone_number or ''}")

    # Địa chỉ
    doc.add_heading("3. Địa chỉ & Quê quán", level=1)
    doc.add_paragraph(f"Địa chỉ thường trú: {app.address or ''}")

    # Người thân
    doc.add_heading("4. Thông tin người thân", level=1)
    doc.add_paragraph(f"Người thân 1: {app.relative1_name or ''} | SĐT: {app.relative1_phone or ''} | Năm sinh: {app.relative1_birthyear or ''}")
    doc.add_paragraph(f"Người thân 2: {app.relative2_name or ''} | SĐT: {app.relative2_phone or ''} | Năm sinh: {app.relative2_birthyear or ''}")

    # Chính sách
    doc.add_heading("5. Chính sách / Hoàn cảnh", level=1)
    doc.add_paragraph(f"Diện chính sách: {app.policy_type or ''}")

    # Ảnh hồ sơ
    doc.add_heading("6. Ảnh hồ sơ", level=1)

    def add_img(title, path):
        doc.add_paragraph(title + ":")
        if not path:
            doc.add_paragraph("Không có")
            return
        img_path = os.path.join("static", path)
        if not os.path.exists(img_path):
            doc.add_paragraph("(File không tồn tại)")
            return
        try:
            doc.add_picture(img_path, width=Inches(2))
        except Exception:
            doc.add_paragraph("(Không thể đọc file ảnh này)")

    add_img("Ảnh 3x4", app.student_photo)
    add_img("Ảnh CCCD", app.citizen_proof)
    add_img("Minh chứng ưu tiên", app.policy_proof)

    # Chữ ký
    doc.add_paragraph("\n\n\n")
    doc.add_paragraph("Người làm đơn", style="Normal").alignment = 2
    doc.add_paragraph(user.fullname or "", style="Normal").alignment = 2

    # Xuất file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"DonDangKy_{user.student_id}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename)

#-------------------------------------------------------
# Quản lý sinh viên
def normalize_text(s):
    """Bỏ dấu và chuyển thành lowercase"""
    if not s:
        return ""
    nfkd_form = unicodedata.normalize('NFKD', s)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)]).lower()

@admin_bp.route("/students")
@login_required
def students():
    keyword = request.args.get("q", "").strip().lower()
    page = request.args.get("page", 1, type=int)
    per_page = 20

    normalized_keyword = normalize_text(keyword)

    students = User.query.filter_by(role="student").all()
    data = []

    for s in students:
        combined = f"{s.student_id} {s.fullname} {s.class_id} {s.citizen_id} {s.phone_number} {s.email or ''}"
        normalized_text = normalize_text(combined)

        if normalized_keyword in normalized_text or not keyword:
            booking = Booking.query.filter_by(user_id=s.id, status="active").first()
            room = Room.query.get(booking.room_id) if booking else None
            payments = Payment.query.filter_by(user_id=s.id).all()
            unpaid = [p for p in payments if p.status != "success"]

            data.append({
                "student": s,
                "room": room,
                "payments": payments,
                "unpaid": unpaid
            })

    total_pages = ceil(len(data) / per_page)
    start = (page - 1) * per_page
    end = start + per_page
    paginated_data = data[start:end]

    # hiển thị dải trang động, chỉ show quanh trang hiện tại
    visible_range = []
    for p in range(1, total_pages + 1):
        if p == 1 or p == total_pages or abs(p - page) <= 2:
            visible_range.append(p)

    return render_template(
        "admin/students.html",
        students=paginated_data,
        keyword=keyword,
        page=page,
        total_pages=total_pages,
        visible_range=visible_range
    )

# Chi tiết sinh viên
@admin_bp.route("/students/<int:student_id>/detail")
@login_required
def student_detail_api(student_id):
    student = User.query.get_or_404(student_id)
    booking = Booking.query.filter_by(user_id=student.id, status="active").first()
    room = Room.query.get(booking.room_id) if booking else None
    payments = Payment.query.filter_by(user_id=student.id).all()

    return jsonify({
        "student_id": student.student_id,
        "fullname": student.fullname,
        "class_id": student.class_id,
        "email": student.email,
        "phone_number": student.phone_number,
        "room": f"{room.room_number} ({room.block})" if room else None,
        "payments": [
            {"service": p.service_type, "amount": p.amount, "status": p.status}
            for p in payments
        ]
    })

#-------------------------------------------------------
# Quản lý hóa đơn
@admin_bp.route("/payments")
@login_required
def payments():
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int, default=datetime.now().year)

    query = Payment.query
    if month:
        query = query.filter(db.extract('month', Payment.created_at) == month)
    if year:
        query = query.filter(db.extract('year', Payment.created_at) == year)

    status = request.args.get("status")           
    service_type = request.args.get("service_type")  
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int)
    min_amount = request.args.get("min_amount", type=float)
    max_amount = request.args.get("max_amount", type=float)
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    # Lọc trạng thái
    if status in ["success", "pending"]:
        query = query.filter(Payment.status == status)

    # Lọc loại dịch vụ
    if service_type:
        query = query.filter(Payment.service_type == service_type)

    # Lọc theo tháng/năm
    if month:
        query = query.filter(Payment.month_paid == month)
    if year:
        query = query.filter(Payment.year_paid == year)

    # Lọc theo khoảng giá trị
    if min_amount:
        query = query.filter(Payment.amount >= min_amount)
    if max_amount:
        query = query.filter(Payment.amount <= max_amount)

    # Lọc theo khoảng thời gian tạo (ngày)
    if date_from:
        try:
            start_date = datetime.strptime(date_from, "%Y-%m-%d")
            query = query.filter(Payment.created_at >= start_date)
        except:
            pass
    if date_to:
        try:
            end_date = datetime.strptime(date_to, "%Y-%m-%d")
            query = query.filter(Payment.created_at <= end_date)
        except:
            pass
    
    payments = query.order_by(Payment.created_at.desc()).all()
    return render_template("admin/payments.html",
                            payments=payments, 
                            current_time=datetime.now(), 
                            filters={
                            "status": status,
                            "service_type": service_type,
                            "month": month,
                            "year": year,
                            "min_amount": min_amount,
                            "max_amount": max_amount,
                            "date_from": date_from,
                            "date_to": date_to
                        })

@admin_bp.route("/payments/<int:payment_id>/mark_paid", methods=["POST"])
@login_required
def mark_payment_paid(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    payment.status = "success"
    db.session.commit()
    flash("Hóa đơn đã được xác nhận thanh toán!", "success")
    return redirect(url_for("admin.payments"))


# Tải file Excel mẫu để nhập hóa đơn
@admin_bp.route('/generate_bills_template')
def generate_bills_template():
    rooms = Room.query.all()

    data = []
    for room in rooms:
        data.append({
            "Room ID": room.id,
            "Address": room.address,
            "Block": room.block,
            "Room Number": room.room_number,
            "Price Room": room.price_room or 0,
            "Electricity Used (kWh)": "",
            "Water Used (m³)": "",
            "Other Fees": "",
            "Note": "",
        })

    df = pd.DataFrame(data)
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)

    now = datetime.now()
    filename = f"bills_template_{now.month}_{now.year}.xlsx"

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# Upload file Excel để tạo hóa đơn từ dữ liệu thực tế
@admin_bp.route('/upload_bills_excel', methods=['POST'])
def upload_bills_excel():
    file = request.files.get('file')
    if not file or not file.filename.endswith('.xlsx'):
        flash('Vui lòng chọn file Excel hợp lệ (.xlsx)', 'danger')
        return redirect(url_for('admin.payments'))

    try:
        df = pd.read_excel(file)
    except Exception as e:
        flash(f'Lỗi khi đọc file Excel: {e}', 'danger')
        return redirect(url_for('admin.payments'))

    created = 0
    skipped = 0
    errors = []
    today = datetime.today()
    current_month, current_year = today.month, today.year

    for index, row in df.iterrows():
        try:
            room_id = int(row['Room ID'])
            room = Room.query.get(room_id)
            if not room:
                skipped += 1
                errors.append(f"Dòng {index+2}: Không tìm thấy phòng ID {room_id}")
                continue

            booking = Booking.query.filter_by(room_id=room_id, status="active").first()
            if not booking:
                skipped += 1
                errors.append(f"Dòng {index+2}: Phòng {room.room_number} chưa có sinh viên đang thuê.")
                continue

            # Đọc dữ liệu từ Excel
            elec_used = float(row.get('Electricity Used (kWh)', 0) or 0)
            water_used = float(row.get('Water Used (m³)', 0) or 0)
            other_fees = float(row.get('Other Fees', 0) or 0)

            elec_cost = elec_used * (room.price_electricity or 0)
            water_cost = water_used * (room.price_water or 0)
            room_cost = float(row.get('Price Room', room.price_room or 0))

            total = elec_cost + water_cost + room_cost + other_fees

            # Tạo hóa đơn mới
            payment = Payment(
                user_id=booking.user_id,
                room_id=room_id,
                service_type='utilities',
                month_paid=current_month,
                year_paid=current_year,
                amount=total,
                payment_method='cash',
                status='pending',
                created_at=today,
                updated_at=today,
            )
            db.session.add(payment)
            created += 1

        except Exception as e:
            skipped += 1
            errors.append(f"Dòng {index+2}: Lỗi không xác định ({e})")

    db.session.commit()

    # Gửi thông báo kết quả
    msg = f"Đã tạo {created} hóa đơn mới."
    if skipped > 0:
        msg += f" Bỏ qua {skipped} dòng lỗi."
    flash(msg, 'success' if created > 0 else 'warning')

    # Ghi log lỗi ra file
    if errors:
        with open("import_bills_log.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(errors))
        flash(f"Chi tiết lỗi được lưu trong file import_bills_log.txt", "info")

    return redirect(url_for('admin.payments'))

#-------------------------------------------------------
# Quản lý yêu cầu dịch vụ
@admin_bp.route("/services")
@login_required
def manage_services():
    status = request.args.get("status")
    service_type = request.args.get("service_type")
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int, default=datetime.now().year)
    min_cost = request.args.get("min_cost", type=float)
    max_cost = request.args.get("max_cost", type=float)
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    query = ServiceRequest.query

    if status:
        query = query.filter(ServiceRequest.status == status)
    if service_type:
        query = query.filter(ServiceRequest.service_type == service_type)
    if month:
        query = query.filter(db.extract("month", ServiceRequest.created_at) == month)
    if year:
        query = query.filter(db.extract("year", ServiceRequest.created_at) == year)
    if min_cost:
        query = query.filter(ServiceRequest.cost >= min_cost)
    if max_cost:
        query = query.filter(ServiceRequest.cost <= max_cost)
    if date_from:
        try:
            start_date = datetime.strptime(date_from, "%Y-%m-%d")
            query = query.filter(ServiceRequest.created_at >= start_date)
        except:
            pass
    if date_to:
        try:
            end_date = datetime.strptime(date_to, "%Y-%m-%d")
            query = query.filter(ServiceRequest.created_at <= end_date)
        except:
            pass

    service_requests = query.order_by(ServiceRequest.created_at.desc()).all()

    return render_template(
        "admin/services.html",
        service_requests=service_requests,
        filters={
            "status": status,
            "service_type": service_type,
            "month": month,
            "year": year,
            "min_cost": min_cost,
            "max_cost": max_cost,
            "date_from": date_from,
            "date_to": date_to,
        },
    )


@admin_bp.route("/services/<int:req_id>/update/<string:status>")
@login_required
def update_service_status(req_id, status):
    req = ServiceRequest.query.get_or_404(req_id)

    if status not in ["in_progress", "completed", "rejected"]:
        flash("Trạng thái không hợp lệ!", "danger")
        return redirect(url_for("admin.manage_services"))

    req.status = status

    # Nếu là dịch vụ thu gom rác → khi completed thì tạo hóa đơn luôn
    if status == "completed" and req.service_type == "trash":
        payment = Payment(
            user_id=req.user.id,
            room_id=req.room.id if req.room else None,
            fullname=req.user.fullname,
            student_id=req.user.student_id,
            class_id=req.user.class_id,
            citizen_id=req.user.citizen_id,
            email=req.user.email,
            phone_number=req.user.phone_number,
            service_type="trash",
            month_paid=db.func.month(db.func.now()),
            year_paid=db.func.year(db.func.now()),
            address=req.room.address if req.room else "-",
            block=req.room.block if req.room else "-",
            room_number=req.room.room_number if req.room else "-",
            amount=req.price,  
            payment_method="cash",
            status="pending"
        )
        db.session.add(payment)
    db.session.commit()

    notif = Notification(
        user_id=req.user.id,
        message=f"Yêu cầu dịch vụ '{req.service_type}' của bạn đã được cập nhật thành '{status}'."
    )
    db.session.add(notif)
    db.session.commit()

    flash("Cập nhật trạng thái thành công!", "success")
    return redirect(url_for("admin.manage_services"))

@admin_bp.route("/services/<int:req_id>/set_price", methods=["POST"])
@login_required
def set_service_price(req_id):
    req = ServiceRequest.query.get_or_404(req_id)

    try:
        price = int(request.form.get("price", 0))
        req.price = price

        # Khi admin nhập giá xong thì cũng tạo Payment
        payment = Payment(
            user_id=req.user.id,
            room_id=req.room.id if req.room else None,
            fullname=req.user.fullname,
            student_id=req.user.student_id,
            class_id=req.user.class_id,
            citizen_id=req.user.citizen_id,
            email=req.user.email,
            phone_number=req.user.phone_number,
            service_type="maintenance",
            month_paid=db.func.month(db.func.now()),
            year_paid=db.func.year(db.func.now()),
            address=req.room.address if req.room else "-",
            block=req.room.block if req.room else "-",
            room_number=req.room.room_number if req.room else "-",
            amount=price,
            payment_method="cash",
            status="pending"
        )
        db.session.add(payment)

        db.session.commit()
        flash("Đã cập nhật giá và tạo hóa đơn!", "success")

        notif = Notification(
            user_id=req.user.id,
            message=f"Dịch vụ '{req.service_type}' của bạn đã được báo giá {price:,}đ. Vui lòng kiểm tra chi tiết."
        )
        db.session.add(notif)
        db.session.commit()

    except Exception:
        db.session.rollback()
        flash("Có lỗi khi cập nhật giá!", "danger")

    return redirect(url_for("admin.manage_services"))

#-------------------------------------------------------
# Quản lý complain
@admin_bp.route("/complains")
@login_required
def manage_complains():
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int, default=datetime.now().year)

    query = Complain.query
    if month:
        query = query.filter(db.extract('month', Complain.created_at) == month)
    if year:
        query = query.filter(db.extract('year', Complain.created_at) == year)
    complains = query.order_by(Complain.created_at.desc()).all()
    return render_template("admin/complains.html", complains=complains, month=month, year=year)

#  Trả lời complain
@admin_bp.route("/complains/<int:complain_id>/reply", methods=["POST"])
@login_required
def reply_complain(complain_id):
    complain = Complain.query.get_or_404(complain_id)
    reply = request.form.get("reply")

    if not reply.strip():
        flash("Nội dung phản hồi không được để trống!", "danger")
        return redirect(url_for("admin.manage_complains"))

    complain.reply = reply
    complain.status = "answered"
    db.session.commit()

    notif = Notification(
        user_id=complain.user_id,
        message=f"Yêu cầu khiếu nại của bạn đã được đóng. Cảm ơn bạn đã phản hồi!"
    )
    db.session.add(notif)
    db.session.commit()


    flash("Đã phản hồi khiếu nại!", "success")
    return redirect(url_for("admin.manage_complains"))

#  Đóng complain
@admin_bp.route("/complains/<int:complain_id>/close", methods=["POST"])
@login_required
def close_complain(complain_id):
    complain = Complain.query.get_or_404(complain_id)
    complain.status = "closed"
    db.session.commit()

    flash("Đã đóng khiếu nại!", "success")
    return redirect(url_for("admin.manage_complains"))

#-------------------------------------------------------
#Announcement
@admin_bp.route("/announcements")
def manage_announcements():
    month = request.args.get("month", type=int)
    year = request.args.get("year", type=int, default=datetime.now().year)
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    query = Announcement.query

    if month:
        query = query.filter(db.extract("month", Announcement.created_at) == month)
    if year:
        query = query.filter(db.extract("year", Announcement.created_at) == year)
    if date_from:
        try:
            start_date = datetime.strptime(date_from, "%Y-%m-%d")
            query = query.filter(Announcement.created_at >= start_date)
        except:
            pass
    if date_to:
        try:
            end_date = datetime.strptime(date_to, "%Y-%m-%d")
            query = query.filter(Announcement.created_at <= end_date)
        except:
            pass

    anns = query.order_by(Announcement.created_at.desc()).all()

    return render_template(
        "admin/announcements.html",
        announcements=anns,
        filters={
            "month": month,
            "year": year,
            "date_from": date_from,
            "date_to": date_to,
        },
    )

@admin_bp.route("/announcements/create", methods=["POST"])
def create_announcement():
    title = request.form.get("title")
    content = request.form.get("content")
    file = request.files.get("image")

    img_url = None
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        save_path = os.path.join(current_app.config["UPLOAD_FOLDER_ANNOUNCEMENTS"], filename)
        file.save(save_path)
        img_url = f"uploads/announcements/{filename}"

    ann = Announcement(title=title, content=content, image_url=img_url, admin_id=current_user.id)
    db.session.add(ann)
    db.session.commit()

    # Tạo thông báo cho tất cả sinh viên
    students = User.query.filter_by(role="student").all()
    for stu in students:
        notif = Notification(
            user_id=stu.id,
            message=f"Thông báo mới: {title}"
        )
        db.session.add(notif)
    db.session.commit()

    flash("Đã đăng thông báo!", "success")
    return redirect(url_for("admin.manage_announcements"))

@admin_bp.route("/announcements/<int:ann_id>")
def announcement_detail(ann_id):
    ann = Announcement.query.get_or_404(ann_id)
    return render_template("admin/announcement_detail.html", ann=ann)

# Sửa Announcement
@admin_bp.route("/announcements/<int:ann_id>/edit", methods=["GET", "POST"])
@login_required
def edit_announcement(ann_id):
    ann = Announcement.query.get_or_404(ann_id)

    if request.method == "POST":
        ann.title = request.form.get("title")
        ann.content = request.form.get("content")

        file = request.files.get("image")
        if file and file.filename != "":
            if file.filename.rsplit(".", 1)[1].lower() in current_app.config["ALLOWED_EXTENSIONS"]:
                filename = secure_filename(file.filename)
                save_path = os.path.join(current_app.config["UPLOAD_FOLDER_ANNOUNCEMENTS"], filename)
                file.save(save_path)
                ann.image_url = f"uploads/announcements/{filename}"

        db.session.commit()
        flash("Cập nhật thông báo thành công!", "success")
        return redirect(url_for("admin.announcement_detail", ann_id=ann.id))

    return render_template("admin/edit_announcement.html", ann=ann)

# Xoá Announcement
@admin_bp.route("/announcements/<int:ann_id>/delete", methods=["POST"])
@login_required
def delete_announcement(ann_id):
    ann = Announcement.query.get_or_404(ann_id)
    db.session.delete(ann)
    db.session.commit()
    flash("Đã xoá thông báo!", "success")
    return redirect(url_for("admin.manage_announcements"))
